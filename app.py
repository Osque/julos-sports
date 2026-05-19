import streamlit as st
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from io import BytesIO
from datetime import datetime
import time

# Imports Selenium
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.common.exceptions import WebDriverException

# --- CONFIGURATION DE LA PAGE ---
st.set_page_config(page_title="Générateur de Calendrier Sportif", layout="wide")

# --- DONNÉES PAR DÉFAUT (MISES À JOUR AVEC FFR / MON CLUB HOUSE) ---
DONNEES_INITIALES = {
    "Rugby": {
        "Stade Langonais (Nationale 2)": "https://monclubhouse.ffr.fr/clubs/stade-langonnais",
        "Salles (Nationale 2)": "https://monclubhouse.ffr.fr/clubs/us-salles",
        "Rugby Club Bassin D'Arcachon (Fédérale 1)": "https://monclubhouse.ffr.fr/clubs/rugby-club-bassin-d-arcachon",
        "Union Athletique Gujan Mestras (Fédérale 1)": "https://monclubhouse.ffr.fr/clubs/u-ath-gujan-mestras",
        "Floirac Rive Droite Rugby (Fédérale 1)": "https://monclubhouse.ffr.fr/clubs/club-municipal-de-floirac",
        "Saint Médard Rugby Club (Fédérale 1)": "https://monclubhouse.ffr.fr/clubs/st-medard-en-jalles-r-c",
        "US Bazas Rugby (Fédérale 2)": "https://monclubhouse.ffr.fr/clubs/us-bazadaise",
        "CA Lormont Rugby (Fédérale 2)": "https://monclubhouse.ffr.fr/clubs/c-a-lormont-hauts-de-garonne",
        "AS Mérignac Rugby (Fédérale 2)": "https://monclubhouse.ffr.fr/clubs/as-merignac-rugby",
        "US Castillonnaise (Fédérale 2)": "https://monclubhouse.ffr.fr/clubs/us-castillonnaise",
        "Stade Bordelais (Fédérale 2)": "https://monclubhouse.ffr.fr/clubs/stade-bordelais",
        "Les Lionnes Elite 1": "https://www.leslionnes-rugby.com/calendrier?team=elite"
    },
    "Football": {
        "SA Gazinet Cestas (National 3)": "https://app.scorenco.com/teams/31817",
        "US Lège Cap Ferret (National 3)": "https://app.scorenco.com/teams/46420",
        "FC BASSIN D'ARCACHON (National 3)": "https://app.scorenco.com/teams/31799",
        "Poule A (Régional 1)": "https://scorenco.com/football/competitions/seniors-regional-1-poule-a-phase-unique-ma8v",
        "Poule B (Régional 1)": "https://scorenco.com/football/competitions/seniors-regional-1-poule-b-phase-unique-orq1"
    },
    "Handball": {
        "CA Béglais (D2 Féminine)": "https://www.ffhandball.fr/competitions/saison-2025-2026-21/national/d2-feminine-2025-26-28228/poule-187259/journee-3/",
        "Stade Pessacais UC (D2 Féminine)": "https://www.ffhandball.fr/competitions/saison-2025-2026-21/national/d2-feminine-2025-26-28228/poule-187259/journee-3/",
        "US Mios Biganos (N1 Féminine)": "https://www.ffhandball.fr/competitions/saison-2025-2026-21/national/nationale-1-feminine-2025-26-28626/poule-169734/journee-26/",
        "Mérignac Gironde HB (N1 Féminine)": "https://www.ffhandball.fr/competitions/saison-2025-2026-21/national/nationale-1-feminine-2025-26-28626/poule-169734/journee-26/"
    }
}

if 'base_clubs' not in st.session_state:
    st.session_state['base_clubs'] = DONNEES_INITIALES

# --- CONFIGURATION SELENIUM ---
@st.cache_resource(show_spinner=False)
def initialiser_navigateur():
    """Configure et lance Chromium en mode silencieux (Headless)"""
    options = Options()
    options.add_argument('--headless')
    options.add_argument('--no-sandbox')
    options.add_argument('--disable-dev-shm-usage')
    options.add_argument('--disable-gpu')
    options.add_argument('--window-size=1920,1080')
    
    try:
        service = Service(ChromeDriverManager().install())
        driver = webdriver.Chrome(service=service, options=options)
        return driver
    except WebDriverException as e:
        st.error(f"Erreur d'initialisation du navigateur : {e}")
        return None

def obtenir_html_rendu(driver, url):
    """Charge la page et attend l'exécution du JavaScript"""
    try:
        driver.get(url)
        time.sleep(5) # Pause nécessaire pour charger les calendriers dynamiques (FFR/FFHandball)
        return driver.page_source
    except Exception as e:
        st.error(f"Erreur lors du chargement de {url} : {e}")
        return None

# --- MOTEUR D'EXTRACTION ---
def analyser_html(html, sport, nom_equipe, url):
    """
    Analyse le code HTML selon qu'il s'agit de la FFR, FFHandball ou Score'n'co.
    """
    matchs_trouves = []
    if not html:
        return matchs_trouves
        
    soup = BeautifulSoup(html, 'html.parser')
    
    # 1. Logique Score'n'co (Classique & App)
    toutes_les_listes = soup.find_all('ul')
    if toutes_les_listes:
        for ul in toutes_les_listes:
            div_date = ul.find('div')
            date_actuelle = div_date.text.strip() if div_date and div_date.text else "Date inconnue"
            
            liens_matchs = ul.find_all('a', href=True)
            for lien in liens_matchs:
                if '/match-' not in lien.get('href', ''):
                    continue
                
                contenu_evenement = lien.find('div', class_='event-content')
                if not contenu_evenement:
                    continue
                
                equipes = contenu_evenement.find_all('div', class_=lambda c: c and 'team' in c.lower())
                equipe_domicile = equipes[0].text.strip() if len(equipes) > 0 else "Inconnu"
                equipe_exterieur = equipes[1].text.strip() if len(equipes) > 1 else "Inconnu"
                
                div_heure = contenu_evenement.find('div', class_='date')
                heure = div_heure.text.strip() if div_heure else "Heure inconnue"
                
                matchs_trouves.append({
                    "Sport": sport, "Équipe Suivie": nom_equipe,
                    "Date": date_actuelle, "Heure": heure,
                    "Rencontre": f"{equipe_domicile} vs {equipe_exterieur}", "Source": url
                })

    # 2. Logique FFR (Mon Club House) & générique
    if not matchs_trouves:
        lignes_tableau = soup.find_all(['tr', 'div']) # Sur FFR, ce sont souvent des flex-divs
        for ligne in lignes_tableau:
            # Cherche des structures texte correspondant à des matchs (ex: "XX/XX/XXXX" ou heures "HH:MM")
            texte_complet = ligne.get_text(separator=' | ', strip=True)
            if " - " in texte_complet and ("202" in texte_complet or ":" in texte_complet): 
                # C'est un filtre heuristique qui capture bien les lignes "Équipe A - Équipe B | 15:00"
                if len(texte_complet) > 15 and len(texte_complet) < 150:
                    matchs_trouves.append({
                        "Sport": sport, "Équipe Suivie": nom_equipe,
                        "Date": "Format Liste", "Heure": "À vérifier",
                        "Rencontre": texte_complet, "Source": url
                    })

    # 3. Échec d'identification HTML
    if not matchs_trouves:
         matchs_trouves.append({
            "Sport": sport, "Équipe Suivie": nom_equipe,
            "Date": "-", "Heure": "-",
            "Rencontre": "Aucun match à venir détecté ou structure à affiner.", "Source": url
        })

    return matchs_trouves

# --- EXPORT WORD ---
def generer_document_word(donnees):
    doc = Document()
    doc.add_heading('Calendrier des Matchs', 0)
    
    if not donnees:
        doc.add_paragraph("Aucun match trouvé pour la période sélectionnée.")
    else:
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Sport'
        hdr_cells[1].text = 'Équipe Suivie'
        hdr_cells[2].text = 'Date & Heure'
        hdr_cells[3].text = 'Rencontre'
        hdr_cells[4].text = 'Lien'
        
        for match in donnees:
            row_cells = table.add_row().cells
            row_cells[0].text = str(match.get('Sport', ''))
            row_cells[1].text = str(match.get('Équipe Suivie', ''))
            row_cells[2].text = f"{match.get('Date', '')} à {match.get('Heure', '')}"
            row_cells[3].text = str(match.get('Rencontre', ''))
            row_cells[4].text = str(match.get('Source', ''))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- INTERFACE UTILISATEUR ---
st.title("📅 Extracteur de Calendrier de Matchs")

onglet_generation, onglet_gestion = st.tabs(["📊 Génération du Planning", "⚙️ Gestion des Équipes et Liens"])

# ==========================================
# ONGLET 1 : GÉNÉRATION
# ==========================================
with onglet_generation:
    col_param, col_res = st.columns([1, 2])
    
    with col_param:
        st.header("Critères")
        date_debut = st.date_input("Date de début", value=datetime.today())
        date_fin = st.date_input("Date de fin", value=datetime.today())
        
        st.subheader("Sélection des équipes")
        equipes_selectionnees = []
        
        for sport, equipes in st.session_state['base_clubs'].items():
            with st.expander(f"🏅 {sport}", expanded=False):
                for nom_equipe, url in equipes.items():
                    coche = st.checkbox(nom_equipe, key=f"check_{sport}_{nom_equipe}")
                    if coche:
                        equipes_selectionnees.append({"Sport": sport, "Nom": nom_equipe, "URL": url})
                        
        lancer_recherche = st.button("Lancer l'extraction", type="primary")

    with col_res:
        st.header("Résultats")
        if lancer_recherche:
            if not equipes_selectionnees:
                st.warning("Veuillez cocher au moins une équipe.")
            else:
                avec_spinner = st.spinner("Analyse des sites officiels via navigateur fantôme...")
                tous_les_matchs = []
                
                with avec_spinner:
                    navigateur = initialiser_navigateur()
                    
                    if navigateur:
                        barre_progression = st.progress(0)
                        total_equipes = len(equipes_selectionnees)
                        
                        for index, eq in enumerate(equipes_selectionnees):
                            st.write(f"Vérification : {eq['Nom']}...")
                            html_rendu = obtenir_html_rendu(navigateur, eq["URL"])
                            matchs_extraits = analyser_html(html_rendu, eq["Sport"], eq["Nom"], eq["URL"])
                            
                            # Nettoyage des doublons liés à la logique générique
                            matchs_uniques = [dict(t) for t in {tuple(d.items()) for d in matchs_extraits}]
                            tous_les_matchs.extend(matchs_uniques)
                            
                            barre_progression.progress((index + 1) / total_equipes)
                            
                        navigateur.quit()
                    else:
                        st.error("Impossible de lancer le module d'extraction web.")
                
                if tous_les_matchs:
                    st.success(f"Extraction terminée !")
                    df = pd.DataFrame(tous_les_matchs)
                    st.dataframe(df, use_container_width=True)
                    
                    fichier_word = generer_document_word(tous_les_matchs)
                    st.download_button(
                        label="📄 Télécharger le planning au format Word",
                        data=fichier_word,
                        file_name="calendrier_matchs.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.info("Aucun match programmé trouvé.")

# ==========================================
# ONGLET 2 : GESTION DES LIENS
# ==========================================
with onglet_gestion:
    st.header("Gérer la base de données des clubs")
    
    col_ajout, col_visu = st.columns([1, 1])
    
    with col_ajout:
        st.subheader("Ajouter une nouvelle équipe")
        nouveau_sport = st.selectbox("Sport", ["Rugby", "Football", "Handball", "Autre"])
        si_autre_sport = st.text_input("Si autre, précisez le sport :")
        sport_final = si_autre_sport if nouveau_sport == "Autre" and si_autre_sport else nouveau_sport
        
        nouveau_nom = st.text_input("Nom de l'équipe")
        nouvelle_url = st.text_input("URL de la page calendrier officielle")
        
        if st.button("Ajouter l'équipe"):
            if sport_final and nouveau_nom and nouvelle_url:
                if sport_final not in st.session_state['base_clubs']:
                    st.session_state['base_clubs'][sport_final] = {}
                st.session_state['base_clubs'][sport_final][nouveau_nom] = nouvelle_url
                st.success(f"L'équipe {nouveau_nom} a été ajoutée !")
                st.rerun()
            else:
                st.error("Veuillez remplir tous les champs.")

    with col_visu:
        st.subheader("Base actuelle")
        for sport, equipes in st.session_state['base_clubs'].items():
            if equipes:
                st.markdown(f"**{sport}**")
                for nom_equipe, url in list(equipes.items()):
                    col_nom, col_btn = st.columns([4, 1])
                    col_nom.text(nom_equipe)
                    if col_btn.button("❌", key=f"del_{sport}_{nom_equipe}", help="Supprimer"):
                        del st.session_state['base_clubs'][sport][nom_equipe]
                        st.rerun()